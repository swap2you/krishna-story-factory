"""Knowledge record package loading, hashing, validation, and export."""
from __future__ import annotations

import hashlib
import json
import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

FIXTURE_MARKER = "TEST FIXTURE — NOT APPROVED DEVOTIONAL CONTENT"
LENSES = ("little_learner", "explorer", "teen", "study")
DEFAULT_LENS = "explorer"


def repo_root() -> Path:
    # apps/api/bhava_api/knowledge/packages.py -> repo root is parents[4]
    return Path(__file__).resolve().parents[4]


def packages_root() -> Path:
    return repo_root() / "content" / "knowledge" / "packages"


def nfc(value: str) -> str:
    return unicodedata.normalize("NFC", value or "")


def sha256_hex(data: str | bytes) -> str:
    if isinstance(data, str):
        data = data.encode("utf-8")
    return hashlib.sha256(data).hexdigest()


def stanza_blocks(content: dict[str, Any]) -> list[dict[str, Any]]:
    blocks = content.get("blocks") or []
    return [b for b in blocks if b.get("block_type") == "stanza"]


def canonical_text_hash(content: dict[str, Any]) -> str:
    parts: list[str] = []
    for block in sorted(stanza_blocks(content), key=lambda b: int(b.get("ord", 0))):
        parts.append(nfc(block.get("devanagari") or ""))
        parts.append(nfc(block.get("iast") or ""))
        parts.append(nfc(block.get("translation_en") or ""))
    return sha256_hex("\n".join(parts))


def scripture_hash(content: dict[str, Any]) -> str:
    parts = [nfc(b.get("devanagari") or "") for b in sorted(stanza_blocks(content), key=lambda b: int(b.get("ord", 0)))]
    return sha256_hex("\n".join(parts))


def translation_hash(content: dict[str, Any]) -> str:
    parts = [nfc(b.get("translation_en") or "") for b in sorted(stanza_blocks(content), key=lambda b: int(b.get("ord", 0)))]
    return sha256_hex("\n".join(parts))


@dataclass
class PackageValidation:
    ok: bool
    errors: list[str]
    package: dict[str, Any] | None = None


def load_package_dir(path: Path) -> dict[str, Any]:
    def read(name: str) -> Any:
        return json.loads((path / name).read_text(encoding="utf-8"))

    return {
        "record": read("record.json"),
        "content": read("content.json"),
        "source_dossier": read("source_dossier.json"),
        "rights": read("rights.json"),
        "assets": read("assets.json"),
        "reviews": read("reviews.json"),
        "manifest": read("manifest.json"),
        "_path": str(path),
    }


def validate_package(pkg: dict[str, Any]) -> PackageValidation:
    errors: list[str] = []
    record = pkg.get("record") or {}
    content = pkg.get("content") or {}
    dossier = pkg.get("source_dossier") or {}

    for key in (
        "record_id",
        "slug",
        "title",
        "content_type",
        "lifecycle",
        "package_status",
        "visibility",
        "source_status",
        "record_version",
        "canonical_text_hash",
        "unicode_normalization",
    ):
        if not record.get(key):
            errors.append(f"record missing {key}")

    if record.get("unicode_normalization") != "NFC":
        errors.append("unicode_normalization must be NFC")

    expected = canonical_text_hash(content)
    if record.get("canonical_text_hash") and record["canonical_text_hash"] != expected:
        errors.append("canonical_text_hash mismatch")

    if record.get("visibility") == "public" and record.get("source_status") != "DOSSIER_READY":
        errors.append("public visibility requires DOSSIER_READY")

    if record.get("source_status") == "SOURCE_BLOCKED" and dossier.get("decision") != "SOURCE_BLOCKED":
        errors.append("source_status SOURCE_BLOCKED requires dossier decision SOURCE_BLOCKED")

    if record.get("fixture"):
        label = record.get("fixture_label") or ""
        if FIXTURE_MARKER not in label:
            errors.append("fixture packages must carry conspicuous fixture_label")
        for block in stanza_blocks(content):
            blob = " ".join(
                [
                    block.get("devanagari") or "",
                    block.get("iast") or "",
                    block.get("translation_en") or "",
                ]
            )
            if "TEST FIXTURE" not in blob:
                errors.append(f"fixture stanza {block.get('block_id')} missing TEST FIXTURE marker")

    # Reject private path leakage markers in public-facing fields
    public_blob = json.dumps({k: pkg[k] for k in ("record", "content", "manifest")}, ensure_ascii=False)
    if re.search(r"(?i)MyPilotDropbox|C:\\\\Users|C:/Users", public_blob):
        errors.append("private operator path leakage detected in package fields")

    return PackageValidation(ok=not errors, errors=errors, package=pkg)


def list_packages() -> list[dict[str, Any]]:
    root = packages_root()
    if not root.exists():
        return []
    out: list[dict[str, Any]] = []
    for child in sorted(root.iterdir()):
        if not child.is_dir():
            continue
        if not (child / "record.json").exists():
            continue
        pkg = load_package_dir(child)
        out.append(pkg)
    return out


def get_package(slug_or_id: str) -> dict[str, Any] | None:
    for pkg in list_packages():
        record = pkg["record"]
        if record.get("slug") == slug_or_id or record.get("record_id") == slug_or_id:
            return pkg
    return None


def is_loopback_host(host: str | None) -> bool:
    if not host:
        return False
    h = host.split("%")[0].lower().strip().strip("[]")
    # Strip :port for IPv4 / hostname; keep IPv6.
    if h.count(":") == 1 and not h.startswith(":"):
        h = h.split(":", 1)[0]
    if h in {"127.0.0.1", "::1", "localhost"}:
        return True
    # Starlette/FastAPI TestClient presents as "testclient".
    if h == "testclient":
        return True
    if h.startswith("127.") and all(part.isdigit() for part in h.split(".")):
        return True
    if h.startswith("::ffff:"):
        mapped = h[7:]
        if mapped.startswith("127.") and all(part.isdigit() for part in mapped.split(".")):
            return True
    return False


VENDORED_FONT_DIR_NAMES = ("assets", "fonts", "noto")
VENDORED_LATIN = "NotoSans-Regular.ttf"
VENDORED_DEVA = "NotoSansDevanagari-Regular.ttf"


def _vendored_font_dirs() -> list[Path]:
    return [
        repo_root().joinpath(*VENDORED_FONT_DIR_NAMES),
        Path("/app").joinpath(*VENDORED_FONT_DIR_NAMES),
        Path(__file__).resolve().parents[4].joinpath(*VENDORED_FONT_DIR_NAMES),
    ]


def _load_expected_font_checksums(font_dir: Path) -> dict[str, str]:
    checksums_path = font_dir / "CHECKSUMS.sha256"
    if not checksums_path.exists():
        raise RuntimeError(f"Missing font checksums at {checksums_path}")
    expected: dict[str, str] = {}
    for line in checksums_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        digest, name = line.split(maxsplit=1)
        expected[name.strip()] = digest.strip()
    return expected


def resolve_vendored_fonts() -> tuple[Path, Path, dict[str, str]]:
    """Return (latin_path, deva_path, sha256_by_filename). Fail closed on mismatch/missing."""
    errors: list[str] = []
    for font_dir in _vendored_font_dirs():
        latin = font_dir / VENDORED_LATIN
        deva = font_dir / VENDORED_DEVA
        if not latin.exists() or not deva.exists():
            errors.append(f"missing fonts under {font_dir}")
            continue
        expected = _load_expected_font_checksums(font_dir)
        for name, path in ((VENDORED_LATIN, latin), (VENDORED_DEVA, deva)):
            if name not in expected:
                raise RuntimeError(f"Checksum missing for {name}")
            actual = sha256_hex(path.read_bytes())
            if actual != expected[name]:
                raise RuntimeError(f"Checksum mismatch for {name}: {actual} != {expected[name]}")
        return latin, deva, {
            VENDORED_LATIN: expected[VENDORED_LATIN],
            VENDORED_DEVA: expected[VENDORED_DEVA],
        }
    raise RuntimeError(
        "Vendored Noto fonts unavailable; PDF export refuses silent glyph fallback. "
        + "; ".join(errors)
    )


def export_manifest(
    pkg: dict[str, Any],
    *,
    format_name: str,
    artifact_sha256: str,
    page_size: str,
    font_hashes: dict[str, str] | None = None,
) -> dict[str, Any]:
    content = pkg["content"]
    record = pkg["record"]
    return {
        "record_id": record["record_id"],
        "record_version": record["record_version"],
        "template_id": f"knowledge-{format_name}-v1",
        "template_version": "1.0.0",
        "scripture_hash": scripture_hash(content),
        "translation_hash": translation_hash(content),
        "canonical_content_hash": canonical_text_hash(content),
        "asset_hashes": [],
        "embedded_font_hashes": font_hashes or {},
        "page_sizes": [page_size],
        "generators": {
            "pdf": "reportlab",
            "docx": "python-docx",
        },
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "validation": {
            "unicode": "pass",
            "study_neutral": True,
            "pdf_ua_claimed": False,
        },
        "artifact_sha256": artifact_sha256,
        "format": format_name,
        "fixture": bool(record.get("fixture")),
    }


def _xml_escape(value: str) -> str:
    return (
        (value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _register_unicode_fonts() -> tuple[str, str, dict[str, str]]:
    """Register vendored Noto fonts only. Fail closed — no system-font fallback."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    latin_path, deva_path, hashes = resolve_vendored_fonts()
    if "BhavaLatin" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("BhavaLatin", str(latin_path)))
    if "BhavaDeva" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("BhavaDeva", str(deva_path)))
    return "BhavaLatin", "BhavaDeva", hashes


def assert_export_allowed(pkg: dict[str, Any]) -> None:
    """Fixture packages may export synthetic text; non-fixture blocked packages may not."""
    record = pkg.get("record") or {}
    rights = pkg.get("rights") or {}
    if record.get("visibility") == "public":
        raise PermissionError("public packages are not exported from private studio routes")
    if record.get("source_status") in {"SOURCE_BLOCKED", "RIGHTS_BLOCKED"} and not record.get("fixture"):
        raise PermissionError("blocked non-fixture packages cannot export scripture bodies")
    if rights.get("download_right") in {"denied", "forbidden", "none"}:
        raise PermissionError("download_right denies export")


def render_pdf(pkg: dict[str, Any], *, page_size: str = "letter") -> tuple[bytes, dict[str, Any]]:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    assert_export_allowed(pkg)
    latin_font, deva_font, font_hashes = _register_unicode_fonts()
    size_key = page_size.lower().strip()
    if size_key not in {"letter", "a4"}:
        raise ValueError("page_size must be letter or a4")
    pagesize = letter if size_key == "letter" else A4

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=pagesize, title=pkg["record"]["title"])
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("KTitle", parent=styles["Heading1"], fontName=latin_font, fontSize=16, leading=20)
    body = ParagraphStyle("KBody", parent=styles["Normal"], fontName=latin_font, fontSize=11, leading=16)
    deva = ParagraphStyle("KDeva", parent=styles["Normal"], fontName=deva_font, fontSize=12, leading=18)
    story = []
    record = pkg["record"]
    story.append(Paragraph(_xml_escape(record["title"]), title_style))
    story.append(
        Paragraph(
            _xml_escape(f"Status: {record.get('source_status')} · {record.get('fixture_label') or ''}"),
            body,
        )
    )
    story.append(Paragraph(_xml_escape(f"Page size: {size_key}"), body))
    story.append(Spacer(1, 12))
    for block in sorted(stanza_blocks(pkg["content"]), key=lambda b: int(b.get("ord", 0))):
        story.append(Paragraph(f"<b>Stanza {block.get('ord')}</b>", body))
        story.append(Paragraph(_xml_escape(nfc(block.get("devanagari") or "")), deva))
        story.append(Paragraph(_xml_escape(nfc(block.get("iast") or "")), body))
        story.append(Paragraph(_xml_escape(nfc(block.get("translation_en") or "")), body))
        story.append(Spacer(1, 10))
    story.append(Paragraph("Export is study-neutral (canonical text only). PDF/UA is not claimed.", body))
    doc.build(story)
    data = buf.getvalue()
    return data, export_manifest(
        pkg,
        format_name="pdf",
        artifact_sha256=sha256_hex(data),
        page_size=size_key,
        font_hashes=font_hashes,
    )


def render_docx(pkg: dict[str, Any], *, page_size: str = "letter") -> tuple[bytes, dict[str, Any]]:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Mm

    assert_export_allowed(pkg)
    size_key = page_size.lower().strip()
    if size_key not in {"letter", "a4"}:
        raise ValueError("page_size must be letter or a4")
    # Fonts are embedded in PDF; for DOCX record the same vendored hashes used by the pilot.
    _, _, font_hashes = resolve_vendored_fonts()

    document = Document()
    section = document.sections[0]
    if size_key == "a4":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    record = pkg["record"]
    document.core_properties.title = record["title"]
    document.core_properties.language = "en"
    h = document.add_heading(record["title"], level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph(f"Status: {record.get('source_status')}")
    if record.get("fixture_label"):
        document.add_paragraph(record["fixture_label"])
    document.add_paragraph(f"Page size: {size_key}")
    for block in sorted(stanza_blocks(pkg["content"]), key=lambda b: int(b.get("ord", 0))):
        document.add_heading(f"Stanza {block.get('ord')}", level=2)
        p = document.add_paragraph(nfc(block.get("devanagari") or ""))
        p.style = document.styles["Normal"]
        document.add_paragraph(nfc(block.get("iast") or ""))
        document.add_paragraph(nfc(block.get("translation_en") or ""))
    document.add_paragraph("Export is study-neutral (canonical text only).")
    buf = BytesIO()
    document.save(buf)
    data = buf.getvalue()
    return data, export_manifest(
        pkg,
        format_name="docx",
        artifact_sha256=sha256_hex(data),
        page_size=size_key,
        font_hashes=font_hashes,
    )